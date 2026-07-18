import json
import os
from docx import Document
from docx.shared import Pt, Inches


def read_text_file(args_json):
    args = json.loads(args_json)
    with open(args["input_path"], "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def create_text_file(args_json):
    args = json.loads(args_json)
    with open(args["output_path"], "w", encoding="utf-8") as f:
        f.write(args.get("content", ""))
    return f"Created text file → {args['output_path']}"


def find_replace_text(args_json):
    args = json.loads(args_json)
    with open(args["input_path"], "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    find_str = args.get("find", "")
    replace_str = args.get("replace", "")
    new_content = content.replace(find_str, replace_str)
    count = content.count(find_str)
    with open(args["output_path"], "w", encoding="utf-8") as f:
        f.write(new_content)
    return f"Replaced {count} occurrences → {args['output_path']}"


def word_count(args_json):
    args = json.loads(args_json)
    with open(args["input_path"], "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    return json.dumps({
        "characters": len(content),
        "characters_no_spaces": len(content.replace(" ", "").replace("\n", "")),
        "words": len(content.split()),
        "lines": content.count("\n") + 1,
        "paragraphs": len([p for p in content.split("\n\n") if p.strip()]),
    })


def read_docx(args_json):
    args = json.loads(args_json)
    doc = Document(args["input_path"])
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return text


def create_docx(args_json):
    args = json.loads(args_json)
    content = json.loads(args.get("content", "{}"))
    doc = Document()
    if "title" in content:
        doc.add_heading(content["title"], 0)
    for para in content.get("paragraphs", []):
        text = para.get("text", "")
        style = para.get("style", "normal")
        if style.startswith("heading"):
            level = int(style.replace("heading", "")) if len(style) > 7 else 1
            doc.add_heading(text, level)
        else:
            doc.add_paragraph(text)
    doc.save(args["output_path"])
    return f"Created DOCX → {args['output_path']}"


def edit_docx(args_json):
    args = json.loads(args_json)
    edits = json.loads(args.get("edits", "[]"))
    doc = Document(args["input_path"])
    for edit in edits:
        if edit.get("action") == "replace":
            find_text = edit.get("find", "")
            replace_text = edit.get("replace", "")
            for para in doc.paragraphs:
                if find_text in para.text:
                    for run in para.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, replace_text)
    doc.save(args["output_path"])
    return f"Edited DOCX → {args['output_path']}"


def docx_to_pdf(args_json):
    args = json.loads(args_json)
    doc = Document(args["input_path"])
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    c = canvas.Canvas(args["output_path"], pagesize=A4)
    w, h = A4
    y = h - 50
    for para in doc.paragraphs:
        if not para.text.strip():
            y -= 15
            continue
        if y < 50:
            c.showPage()
            y = h - 50
        font = "Helvetica-Bold" if para.style.name.startswith("Heading") else "Helvetica"
        size = 14 if para.style.name.startswith("Heading") else 11
        c.setFont(font, size)
        c.drawString(50, y, para.text[:100])
        y -= size + 6
    c.save()
    return f"Converted DOCX → PDF: {args['output_path']}"
